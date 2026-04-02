import io
from pathlib import Path
from typing import List, Tuple

from PyPDF2 import PdfReader
from docx import Document


def extract_text_from_pdf(file_bytes: bytes, filename: str) -> str:
    """Extract all text from a PDF file."""
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                pages.append(text.strip())

        if not pages:
            return f"[WARNING: PDF '{filename}' contained no extractable text (possibly scanned/image-based)]"

        return "\n\n".join(pages)
    except Exception as e:
        return f"[ERROR: Could not read PDF '{filename}': {str(e)}]"


def extract_text_from_docx(file_bytes: bytes, filename: str) -> str:
    """Extract all text from a DOCX file."""
    try:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        if not paragraphs:
            return f"[WARNING: DOCX '{filename}' contained no text]"

        return "\n\n".join(paragraphs)
    except Exception as e:
        return f"[ERROR: Could not read DOCX '{filename}': {str(e)}]"


def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """Route to the appropriate extractor based on file extension."""
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes, filename)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(file_bytes, filename)
    elif ext in (".txt", ".md", ".csv"):
        # Plain text files — decode directly
        try:
            return file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            return file_bytes.decode("latin-1")
    else:
        return f"[UNSUPPORTED: File '{filename}' has unsupported extension '{ext}'. Supported: .pdf, .docx, .doc, .txt]"


def combine_texts(pasted_text: str, file_results: List[Tuple[str, str]]) -> str:
    """
    Combine pasted text and extracted file text into a single document.

    Args:
        pasted_text: Text pasted directly by the user
        file_results: List of (filename, extracted_text) tuples

    Returns:
        Combined text with clear section separators
    """
    sections = []

    if pasted_text and pasted_text.strip():
        sections.append("=== PASTED TEXT ===\n" + pasted_text.strip())

    for filename, text in file_results:
        sections.append(f"=== FILE: {filename} ===\n" + text.strip())

    return "\n\n".join(sections)
